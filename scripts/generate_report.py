"""
JobForge AI — Project Report Generator.

Generates a formatted .docx (and optionally PDF) project report.

Usage:
    pip install python-docx
    python scripts/generate_report.py
    python scripts/generate_report.py --pdf   # also compile to PDF via pdflatex

Output: outputs/JobForge_Report_2026-03-12.docx
"""

from __future__ import annotations

import argparse
import subprocess
import sys
from datetime import date
from pathlib import Path

# ── Auto-install python-docx if missing ──────────────────────────────────────
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

# ── Colours ───────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x2A, 0x4A)
TEAL   = RGBColor(0x00, 0x7A, 0x87)
DARK   = RGBColor(0x22, 0x22, 0x22)
GREY   = RGBColor(0x55, 0x55, 0x55)
GREEN  = RGBColor(0x27, 0x7A, 0x3A)
RED    = RGBColor(0xC0, 0x39, 0x2B)
AMBER  = RGBColor(0xE6, 0x7E, 0x22)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

TODAY  = date.today().isoformat()  # 2026-03-12
OUTPUT = Path(__file__).parent.parent / "outputs" / f"JobForge_Report_{TODAY}.docx"


# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_colour: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def _add_run(para, text: str, bold=False, italic=False,
             colour: RGBColor | None = None, size: int | None = None) -> None:
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if colour:
        run.font.color.rgb = colour
    if size:
        run.font.size = Pt(size)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY if level == 1 else TEAL
        run.font.size = Pt(18 if level == 1 else 14 if level == 2 else 12)


def _para(doc: Document, text: str = "", bold=False, italic=False,
          colour: RGBColor | None = None, size: int = 10) -> None:
    p = doc.add_paragraph()
    _add_run(p, text, bold=bold, italic=italic,
             colour=colour or DARK, size=size)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(0)


def _bullet(doc: Document, text: str, bold_prefix: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        _add_run(p, bold_prefix, bold=True, colour=DARK, size=10)
    _add_run(p, text, colour=DARK, size=10)
    p.paragraph_format.space_after = Pt(2)


def _code(doc: Document, text: str) -> None:
    """Mono-spaced block for code/config snippets."""
    p  = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name    = "Courier New"
    run.font.size    = Pt(8.5)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)


def _table(doc: Document, headers: list[str], rows: list[list[str]],
           col_widths: list[float] | None = None) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        _set_cell_bg(cell, "1B2A4A")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                run.font.color.rgb = DARK
            if r_idx % 2 == 1:
                _set_cell_bg(cell, "EFF3FB")

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()


# ── Cover Page ────────────────────────────────────────────────────────────────

def _cover(doc: Document) -> None:
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(title, "JobForge AI", bold=True, colour=NAVY, size=32)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(subtitle, "Autonomous Multi-Agent Job Hunting Pipeline", colour=TEAL, size=16)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(meta, f"Project Report  ·  {TODAY}  ·  Status: Live on Railway",
             colour=GREY, size=10, italic=True)

    doc.add_paragraph()
    doc.add_paragraph()

    by = doc.add_paragraph()
    by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(by, "Viraj Bulugahapitiya", bold=True, colour=DARK, size=11)

    doc.add_page_break()


# ── Section builders ──────────────────────────────────────────────────────────

def _section_overview(doc: Document) -> None:
    _heading(doc, "1. Project Overview")
    _para(doc,
          "JobForge AI is an autonomous multi-agent job hunting pipeline built for UK "
          "AI / ML / Data Science roles. Given a master CV and a structured Skill Inventory, "
          "it scrapes thousands of job listings every Wednesday, ML-filters and LLM-scores "
          "each listing against the candidate profile, generates tailored LaTeX CVs compiled "
          "to PDF for the top matches, and emails a formatted Excel digest with attached CVs — "
          "all with zero manual input per run.")


def _section_architecture(doc: Document) -> None:
    _heading(doc, "2. Architecture")

    _heading(doc, "2.1 Pipeline DAG", level=2)
    _code(doc, "Scout ──► Matchmaker ──► Tailor ──► Dispatcher ──► END")
    _code(doc, "          (0 qualified) ─────────────────────────────────►")
    _code(doc, "(0 new jobs) ────────────────────────────────────────────►")
    _para(doc, "Implemented as a LangGraph StateGraph with conditional edges that short-circuit "
               "to Dispatcher if no new jobs or no qualified jobs are found.")

    _heading(doc, "2.2 Agent Summary", level=2)
    _table(doc,
           ["Agent", "Type", "LLM", "Responsibility"],
           [
               ["Scout",      "Deep Agent (plan/execute/reflect)", "None",         "Scrape + cross-run deduplicate jobs"],
               ["Matchmaker", "Deep Agent",                        "Gemini Flash",  "ML pre-screen + LLM score all jobs"],
               ["Tailor",     "Deep Agent",                        "Gemini Pro",    "Generate job-specific LaTeX + PDF CVs"],
               ["Dispatcher", "Deterministic",                     "None",          "Excel digest + SMTP / Resend email"],
           ],
           col_widths=[3.0, 4.5, 3.0, 6.5])

    _heading(doc, "2.3 Data Flow per Run", level=2)
    _table(doc,
           ["Stage", "Count (live run)", "Mechanism"],
           [
               ["Raw scrape",        "577",  "6 connectors: Adzuna, Reed, Wellfound, LinkedIn proxy, Indeed proxy, Career pages"],
               ["After dedup",       "315",  "DedupStore — SQLite hash of (title + company + location)"],
               ["After ML gate",     "~130", "SBERT cosine + BM25 + Jaccard skill overlap ensemble ≥ 0.30"],
               ["After score cache", "~90",  "ScoreCache (14-day TTL) — skip LLM if same job seen recently"],
               ["LLM scored",        "234",  "Gemini Flash ×15 parallel with Semaphore rate limit"],
               ["Qualified",         "171",  "overall_score ≥ 70 (configurable threshold)"],
               ["CVs tailored",      "20",   "Top-20 by score — Gemini Pro ×5 parallel"],
               ["Email sent",        "1",    "Excel + top-10 PDFs attached, SMTP dispatch"],
           ],
           col_widths=[4.0, 3.5, 9.5])

    _heading(doc, "2.4 Tech Stack", level=2)
    _table(doc,
           ["Layer", "Technology", "Purpose"],
           [
               ["Agent orchestration", "LangGraph StateGraph",              "DAG, conditional edges, shared state"],
               ["LLM — scoring",       "Gemini 2.5 Flash",                  "Cost-effective bulk scoring (~$0.075/1M tokens)"],
               ["LLM — CV writing",    "Gemini 2.5 Pro",                    "High-quality LaTeX rewriting"],
               ["LLM framework",       "LangChain",                         "ChatGoogleGenerativeAI, message types"],
               ["Dense retrieval",     "sentence-transformers (MiniLM-L6)", "Semantic similarity pre-screen"],
               ["Sparse retrieval",    "rank-bm25 (Okapi BM25)",            "Keyword / exact match pre-screen"],
               ["Validation",          "Pydantic V2",                       "LLM JSON output schema enforcement"],
               ["Settings",            "pydantic-settings",                 "Env var loading, Railway-compatible"],
               ["Persistence",         "SQLite (local) / PostgreSQL (prod)","Dedup, score cache, run history"],
               ["CV compilation",      "pdflatex (texlive) / tectonic",     "LaTeX → PDF"],
               ["Excel output",        "pandas + openpyxl",                 "Colour-coded digest"],
               ["Email",               "SMTP / Resend",                     "Weekly digest dispatch"],
               ["Deployment",          "Railway cron service",              "Weekly autonomous execution"],
               ["Logging",             "structlog (JSON)",                  "Structured, Railway-compatible"],
           ],
           col_widths=[4.0, 4.5, 8.5])


def _section_phases(doc: Document) -> None:
    _heading(doc, "3. What Was Built — Phase by Phase")

    _heading(doc, "Phase 1 — Infrastructure", level=2)
    _bullet(doc, "6 connectors: Adzuna, Reed, Wellfound, LinkedIn proxy, Indeed proxy, Career pages")
    _bullet(doc, "Pydantic V2 models: RawJob, ScoredJob, MatchScore, SkillInventory, TailoredCV, JobForgeState")
    _bullet(doc, "DedupStore (cross-run job dedup by content hash) + RunHistory (telemetry) — both SQLite")
    _bullet(doc, "Scout Agent: plan/execute/reflect/output lifecycle, queries all sources, deduplicates")
    _bullet(doc, "Dispatcher Agent: colour-coded Excel with openpyxl, SMTP email with attachments")
    _bullet(doc, "Structured prompts (system + user templates) for Matchmaker and Tailor")

    _heading(doc, "Phase 2 — Matchmaker LLM Integration", level=2)
    _bullet(doc, "_score_job() calls ChatGoogleGenerativeAI (Gemini 2.5 Flash)")
    _bullet(doc, "Scores 6 dimensions: technical_skills, domain_experience, seniority_fit, location, visa, role_alignment")
    _bullet(doc, "Pydantic V2 validates JSON response — rejects malformed LLM output")
    _bullet(doc, "PSW Graduate Route visa logic: +10pts sponsoring roles, −5pts UK-citizens-only")

    _heading(doc, "Phase 3 — Tailor Agent LLM Integration", level=2)
    _bullet(doc, "_modify_latex() calls Gemini Pro with full master CV + skill inventory + projects bank")
    _bullet(doc, "Rewrites Professional Summary, reorders Technical Skills, selects 3–4 best projects")
    _bullet(doc, "Hallucination detector: scans modified LaTeX for metrics/skills not in SkillInventory")
    _bullet(doc, "_compile_pdf(): pdflatex primary, tectonic fallback")
    _bullet(doc, "3 CV variants: ai_engineer, data_scientist, ml_engineer")

    _heading(doc, "Phase 4 — Deployment to Railway", level=2)
    _bullet(doc, "Dockerfile: python:3.11-slim + texlive-latex-extra + texlive-fonts-extra + tectonic fallback")
    _bullet(doc, "railway.toml: single cron service, cronSchedule = '0 7 * * 3' (Wednesday 07:00 UTC)")
    _bullet(doc, "DATABASE_URL env var auto-provided by Railway PostgreSQL addon")
    _bullet(doc, "settings.py normalises postgresql:// → postgresql+asyncpg:// for SQLAlchemy async")


def _section_bugs(doc: Document) -> None:
    _heading(doc, "4. What Went Wrong & How We Fixed It")

    bugs = [
        (
            "Bug 1 — CVs shorter than master template",
            "base_latex[:5000] in _modify_latex() truncated the CV at ~5,000 chars. "
            "Education and Projects sections were cut off. The LLM only regenerated what it saw.",
            "Removed the [:5000] slice. Added explicit prompt instruction: "
            "'Return the COMPLETE LaTeX document. Do NOT truncate or omit any sections.' "
            "Gemini Pro's 1M-token context window handles the full CV easily."
        ),
        (
            "Bug 2 — No PDFs, only .tex files",
            "pdflatex not installed on local Windows machine or Railway container. "
            "_compile_pdf() caught FileNotFoundError silently and returned False.",
            "Local: install MiKTeX on Windows. "
            "Railway: added texlive packages to Dockerfile covering fontawesome5 dependency. "
            "Added tectonic as a second fallback compiler tried after pdflatex."
        ),
        (
            "Bug 3 — Railway build: OSError: Readme file does not exist",
            "pyproject.toml declared readme = 'README.md' but no README.md existed. "
            "Hatchling validates this at build time and hard-crashes.",
            "Removed the readme field from pyproject.toml."
        ),
        (
            "Bug 4 — Railway build: editable install failed",
            "Dockerfile had COPY pyproject.toml then pip install -e . immediately, "
            "but src/jobforge/ source code was not yet present for editable install.",
            "Moved COPY . . before pip install . "
            "Switched from -e . (editable) to . (regular install). "
            "Editable mode is meaningless in a container."
        ),
        (
            "Bug 5 — 25+ minute pipeline runtime",
            "Matchmaker made 234 LLM calls sequentially (~3–5s each = 12–20 min). "
            "Tailor made 165 Pro calls sequentially.",
            "Parallelised both agents with asyncio.gather + Semaphore. "
            "Added ML pre-screen gate and CV cap. See Section 5."
        ),
    ]

    for title, cause, fix in bugs:
        p = doc.add_paragraph()
        _add_run(p, title, bold=True, colour=NAVY, size=10)
        _para(doc, f"Cause: {cause}", italic=True, colour=GREY)
        _para(doc, f"Fix:     {fix}", colour=DARK)
        doc.add_paragraph()


def _section_cost(doc: Document) -> None:
    _heading(doc, "5. Cost & Performance Optimisations")

    _heading(doc, "5.1 Sequential → Parallel Execution", level=2)
    _para(doc, "Both Matchmaker and Tailor were refactored from sequential for-loops to "
               "asyncio.gather() with Semaphore rate limiting.")
    _table(doc,
           ["Agent", "Concurrency", "Before", "After", "Reason for limit"],
           [
               ["Matchmaker", "Semaphore(15)", "~12 min", "~47 sec", "Gemini Flash RPM rate limit"],
               ["Tailor",     "Semaphore(5)",  "~8 min",  "~2 min",  "Gemini Pro slower + costlier"],
           ],
           col_widths=[3.0, 3.5, 2.5, 2.5, 5.5])

    _heading(doc, "5.2 Shared LLM Instance", level=2)
    _para(doc, "Previously a new ChatGoogleGenerativeAI object (and HTTP client) was instantiated "
               "per job — creating 234 objects per run. Now created once at the start of execute() "
               "and shared across all concurrent coroutines.")

    _heading(doc, "5.3 Score Cache (SQLite, 14-day TTL)", level=2)
    _para(doc, "New score_cache table stores the full Gemini Flash result keyed on job.dedup_hash. "
               "Any time the same job reappears within 14 days, the cached score is returned instantly "
               "at zero LLM cost. TTL is 14 days (2× the weekly cadence) so entries survive to the "
               "next Wednesday run. Estimated steady-state cache hit rate: ~60% on re-posted UK listings.")
    _table(doc,
           ["Run", "LLM calls", "Notes"],
           [
               ["Run 1 (cold)", "~130", "All jobs new, no cache"],
               ["Run 2 (warm)", "~50",  "~60% cache hit on re-posted jobs"],
               ["Run N (steady)", "~20–30", "Most active listings are cached"],
           ],
           col_widths=[4.0, 3.5, 9.5])

    _heading(doc, "5.4 ML Pre-screen Gate — Three-Signal Ensemble", level=2)
    _para(doc, "New module: src/jobforge/ml/prescreen.py. Runs entirely on CPU before any LLM call. "
               "Filters jobs with no meaningful overlap with the skill profile using three complementary signals "
               "combined as a weighted ensemble (0.50 · dense + 0.30 · BM25 + 0.20 · Jaccard ≥ 0.30).")
    _table(doc,
           ["Signal", "Technique", "Library", "What it catches"],
           [
               ["Dense",  "SBERT cosine similarity",  "sentence-transformers (all-MiniLM-L6-v2, 22 MB)", "'built ML pipelines' ≈ 'developed AI systems'"],
               ["Sparse", "BM25 / Okapi BM25",         "rank-bm25",                                        "Exact matches: 'PyTorch', 'LangGraph', 'FastAPI'"],
               ["Exact",  "Jaccard skill overlap",     "built-in set ops",                                  "Raw inventory skill token intersection"],
           ],
           col_widths=[2.0, 4.0, 5.5, 5.5])
    _para(doc, "SBERT model is pre-downloaded during Docker image build (cached in image layer) to avoid "
               "cold-start download on Railway.")

    _heading(doc, "5.5 CV Cap — Top 20 per Run", level=2)
    _para(doc, "All qualified jobs appear in the Excel digest. Only the top-20 by overall_score "
               "receive tailored CV generation. Configurable via MAX_CVS_PER_RUN env var on Railway.")
    _table(doc,
           ["Metric", "Before", "After"],
           [
               ["Tailor LLM calls / run",  "165",        "20"],
               ["Matchmaker calls / run",  "315 (seq.)", "~50 (parallel + cache + ML gate)"],
               ["Total runtime",           "~25–30 min", "~5–8 min"],
               ["Estimated weekly cost",   "~$1.50",     "~$0.15–$0.25"],
           ],
           col_widths=[6.0, 4.0, 4.0])


def _section_deployment(doc: Document) -> None:
    _heading(doc, "6. Infrastructure & Deployment")

    _heading(doc, "6.1 Local Development", level=2)
    _bullet(doc, "Python 3.11 virtual environment (.venv)")
    _bullet(doc, "SQLite database at data/jobforge.db")
    _bullet(doc, "pdflatex via MiKTeX (Windows installation required)")
    _bullet(doc, ".env file with all API keys (from .env.example)")

    _heading(doc, "6.2 Railway Production", level=2)
    _bullet(doc, "1 service only — cron type, no web server needed")
    _bullet(doc, "Schedule: 0 7 * * 3 — every Wednesday 07:00 UTC")
    _bullet(doc, "Database: Railway PostgreSQL addon (auto-provides DATABASE_URL)")
    _bullet(doc, "Dockerfile: python:3.11-slim + full texlive + SBERT model pre-cached")

    _heading(doc, "6.3 Key Environment Variables", level=2)
    _table(doc,
           ["Variable", "Value / Source", "Notes"],
           [
               ["GEMINI_API_KEY",          "Google AI Studio",        "Used by both Flash and Pro models"],
               ["DATABASE_URL",            "Auto-set by Railway",     "postgres:// normalised to postgresql+asyncpg://"],
               ["ADZUNA_APP_ID/KEY",       "Adzuna developer portal", "Primary UK job source"],
               ["REED_API_KEY",            "Reed developer portal",   "Second primary UK source"],
               ["TAVILY_API_KEY",          "Tavily",                  "Web search for career pages"],
               ["SMTP_USER/PASSWORD",      "Gmail app password",      "Or use RESEND_API_KEY"],
               ["RECIPIENT_EMAIL",         "Your email",              "Weekly digest destination"],
               ["MAX_CVS_PER_RUN",         "20 (default)",            "Cost control — top-N CV generation"],
               ["ML_PRESCREEN_THRESHOLD",  "0.30 (default)",          "Ensemble gate cutoff"],
               ["MATCH_THRESHOLD",         "70 (default)",            "LLM score cutoff for qualifying"],
           ],
           col_widths=[4.5, 4.0, 8.5])


def _section_results(doc: Document) -> None:
    _heading(doc, "7. Live Run Results — 2026-03-11")

    _para(doc, "First successful autonomous run on Railway:")
    _table(doc,
           ["Metric", "Value"],
           [
               ["Total scraped",      "577"],
               ["After dedup",        "315"],
               ["LLM scored",         "234"],
               ["Qualified (≥65%)",   "171"],
               ["Sponsoring roles",   "0"],
               ["Startup roles",      "60"],
               ["Score 90–100%",      "59 jobs"],
               ["Score 80–89%",       "68 jobs"],
               ["Score 70–79%",       "37 jobs"],
               ["Top match",          "96% at Unknown Startup"],
               ["CVs generated",      "165 (pre-cap run)"],
           ],
           col_widths=[6.0, 11.0])


def _section_roadmap(doc: Document) -> None:
    _heading(doc, "8. Future Roadmap")

    _heading(doc, "Near-term (can build now)", level=2)

    _bullet(doc, " Learned Calibration Gate — after 5–10 weekly runs, "
                 "train LogisticRegression on [embedding_score, bm25_score, skill_overlap, "
                 "is_startup, offers_sponsorship, salary_band] to predict LLM_score ≥ 70. "
                 "Replaces hand-tuned threshold with a self-calibrating model.",
            bold_prefix="A.")

    _bullet(doc, " Salary NER — regex extraction of '£45k–£65k' patterns from job body text "
                 "when the salary field is empty. Fills missing data for better scoring context.",
            bold_prefix="B.")

    _bullet(doc, " HTML email — replace plain-text body with styled HTML: "
                 "role-by-role summary, clickable Apply URLs, colour-coded score bands.",
            bold_prefix="C.")

    _heading(doc, "Medium-term", level=2)

    _bullet(doc, " Cross-encoder re-ranking — replace 'top-20 by LLM score' with "
                 "cross-encoder/ms-marco-MiniLM-L-6-v2. Cross-encoders attend to query + document "
                 "together (full attention) — more accurate for final CV selection.",
            bold_prefix="D.")

    _bullet(doc, " MinHash LSH deduplication — current dedup is exact hash. MinHash estimates "
                 "Jaccard on character n-gram shingles, catching 'Senior ML Engineer' ≈ "
                 "'Sr. Machine Learning Engineer' at the same company.",
            bold_prefix="E.")

    _bullet(doc, " Cover letter generation — 5th agent, short targeted cover letter per role "
                 "using same Tailor prompt pattern. 2–3 sentences, no hallucinations.",
            bold_prefix="F.")

    _heading(doc, "Long-term", level=2)

    _bullet(doc, " Application tracking — store applied/not-applied status per job. "
                 "Second digest email with follow-up reminders 7 days after application.",
            bold_prefix="G.")

    _bullet(doc, " Dashboard — FastAPI + HTML dashboard showing run history, score trends, "
                 "application funnel. Single Railway web service alongside the cron.",
            bold_prefix="H.")


def _section_ml(doc: Document) -> None:
    _heading(doc, "9. ML / DL Techniques Used & Planned")

    _table(doc,
           ["Technique", "Status", "Where", "Interview Topics"],
           [
               ["SBERT bi-encoder",        "Live",    "ML pre-screen Signal 1",    "Contrastive learning, NLI fine-tuning, sentence embeddings, ANN search"],
               ["BM25 / Okapi BM25",       "Live",    "ML pre-screen Signal 2",    "TF-IDF, k1/b params, Elasticsearch, sparse vs dense retrieval"],
               ["Hybrid search",           "Live",    "Ensemble of 3 signals",     "Dense + sparse = SOTA for RAG, production IR systems"],
               ["Jaccard similarity",      "Live",    "ML pre-screen Signal 3",    "Set similarity, precision vs recall trade-off, MinHash"],
               ["Cosine similarity",       "Live",    "SBERT scoring",             "L2 normalisation trick, dot product = cosine when normalised"],
               ["LLM structured output",   "Live",    "Matchmaker + Tailor",       "Prompt engineering, JSON schema validation, Pydantic"],
               ["Score cache (TTL)",       "Live",    "ScoreCache SQLite",         "Memoisation, cache invalidation strategies, TTL design"],
               ["Async concurrency",       "Live",    "Both agents",               "I/O-bound parallelism, Semaphore, asyncio.gather"],
               ["Logistic Regression",     "Planned", "Learned calibration gate",  "Binary classification, feature engineering, temporal split"],
               ["Cross-encoder",           "Planned", "Tailor selection",          "Bi-encoder vs cross-encoder, ColBERT, re-ranking pipelines"],
               ["MinHash LSH",             "Planned", "Near-dedup detection",      "Locality-sensitive hashing, Jaccard estimation, ANN"],
               ["Salary NER",              "Planned", "JD parsing",                "NER, IOB tagging, rule-based vs learned NER, regex + ML hybrid"],
           ],
           col_widths=[3.8, 2.0, 4.2, 7.0])


# ── Main ──────────────────────────────────────────────────────────────────────

def build_report() -> Path:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default font
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    _cover(doc)
    _section_overview(doc)
    doc.add_page_break()
    _section_architecture(doc)
    doc.add_page_break()
    _section_phases(doc)
    doc.add_page_break()
    _section_bugs(doc)
    doc.add_page_break()
    _section_cost(doc)
    doc.add_page_break()
    _section_deployment(doc)
    doc.add_page_break()
    _section_results(doc)
    _section_roadmap(doc)
    doc.add_page_break()
    _section_ml(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    return OUTPUT


def compile_pdf(docx_path: Path) -> Path | None:
    """
    Attempt to convert .docx → PDF via LibreOffice headless (cross-platform).
    Falls back gracefully if LibreOffice is not installed.
    """
    for binary in ["soffice", "libreoffice"]:
        try:
            subprocess.run(
                [binary, "--headless", "--convert-to", "pdf",
                 "--outdir", str(docx_path.parent), str(docx_path)],
                check=True, capture_output=True, timeout=60,
            )
            pdf = docx_path.with_suffix(".pdf")
            if pdf.exists():
                return pdf
        except (FileNotFoundError, subprocess.CalledProcessError):
            continue
    return None


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate JobForge project report")
    parser.add_argument("--pdf", action="store_true",
                        help="Also convert to PDF via LibreOffice headless")
    args = parser.parse_args()

    print("Building report...")
    docx_path = build_report()
    print(f"  [OK] DOCX saved: {docx_path}")

    if args.pdf:
        print("Converting to PDF (requires LibreOffice)...")
        pdf_path = compile_pdf(docx_path)
        if pdf_path:
            print(f"  [OK] PDF saved:  {pdf_path}")
        else:
            print("  [!]  LibreOffice not found -- open the .docx in Word and Save As PDF")
